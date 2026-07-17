"""
docx_export.py — Generates a resume/portfolio-summary Word document using python-docx.
Mirrors pdf_export.py so visitors can choose their preferred format.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import projects_data
from projects_data import CATEGORY_LABELS

PRIMARY_COLOR = RGBColor(0x00, 0xA3, 0xFF)
ACCENT_COLOR = RGBColor(0x7B, 0x00, 0xFF)


def build_resume_docx():
    doc = Document()

    title = doc.add_heading("Ravi Der", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph("AI / Machine Learning & Python Developer")
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_heading("About", level=1)
    doc.add_paragraph(
        "B.Tech in Information Technology, currently pursuing an M.Tech in Artificial "
        "Intelligence, Machine Learning and Data Science. Based in Mandvi, District "
        "Bhavnagar, India. Builder of 20+ AI/ML, data analysis and Python/Django "
        "projects, with 3 completed internships."
    )

    doc.add_heading("Core Skills", level=1)
    skills_table_data = [
        ("Backend", "Python, FastAPI, MySQL, REST APIs, JWT Auth"),
        ("AI / ML", "Scikit-Learn, Pandas, NumPy, Matplotlib, Feature Engineering"),
        ("Frontend", "HTML5, CSS3, JavaScript (ES6+), Bootstrap 5"),
        ("DevOps", "Git, GitHub Actions, CLI/Bash, Docker (containerization)"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in skills_table_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_heading("Selected Projects", level=1)
    by_category = {}
    for p in projects_data.get_projects():
        by_category.setdefault(p["category"], []).append(p)

    for cat, items in by_category.items():
        doc.add_heading(CATEGORY_LABELS.get(cat, cat), level=2)
        for p in items:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(p["title"])
            run.bold = True
            para.add_run(f" — {', '.join(p['tags'])}")

    doc.add_heading("Contact", level=1)
    doc.add_paragraph(
        "Full project details, live demos and a contact form are available on the "
        "portfolio website. GitHub: github.com/derravi"
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
